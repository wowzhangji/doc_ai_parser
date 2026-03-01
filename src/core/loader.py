"""文档加载模块 - 支持PDF、Word、Excel等多种格式"""

import io
import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import fitz  # pymupdf
import pdfplumber
import pandas as pd
from docx import Document as DocxDocument

from src.models.document import Document, DocumentSection, DocumentStructure, SectionType


class DocumentLoader(ABC):
    """文档加载器基类"""
    
    @abstractmethod
    def load(self, source: Union[Path, bytes, str]) -> Document:
        """加载文档并返回结构化内容"""
        pass
    
    def _detect_api_sections(self, text: str) -> List[DocumentSection]:
        """检测API相关章节"""
        sections = []
        
        # API端点模式匹配
        api_patterns = [
            r'(GET|POST|PUT|DELETE|PATCH)\s+(/[\w/{}-]+)',  # REST API
            r'(https?://[^\s]+)',  # URL
            r'Endpoint:\s*([^\n]+)',  # Endpoint声明
        ]
        
        lines = text.split('\n')
        current_section = None
        
        for line in lines:
            # 检测标题
            if line.strip().startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title = line.strip('# ').strip()
                sections.append(DocumentSection(
                    type=SectionType.HEADING,
                    content=title,
                    level=level
                ))
            # 检测API端点
            elif any(re.search(pattern, line) for pattern in api_patterns):
                sections.append(DocumentSection(
                    type=SectionType.API_ENDPOINT,
                    content=line.strip(),
                    metadata={"is_endpoint": True}
                ))
            # 检测代码块
            elif line.strip().startswith('```'):
                if current_section and current_section.type == SectionType.CODE:
                    sections.append(current_section)
                    current_section = None
                else:
                    current_section = DocumentSection(
                        type=SectionType.CODE,
                        content="",
                        metadata={"language": line.strip()[3:]}
                    )
            elif current_section and current_section.type == SectionType.CODE:
                current_section.content += line + '\n'
            else:
                # 普通段落
                if line.strip():
                    sections.append(DocumentSection(
                        type=SectionType.PARAGRAPH,
                        content=line.strip()
                    ))
        
        return sections


class PDFLoader(DocumentLoader):
    """PDF文档加载器"""
    
    def load(self, source: Union[Path, bytes, str]) -> Document:
        """加载PDF文档"""
        if isinstance(source, (str, Path)):
            file_path = Path(source)
            file_content = None
            doc = fitz.open(file_path)
        else:
            file_path = None
            file_content = source
            doc = fitz.open(stream=source, filetype="pdf")
        
        try:
            full_text = ""
            metadata = {
                "page_count": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
            }
            
            # 提取文本
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                full_text += f"\n\n--- Page {page_num + 1} ---\n\n"
                full_text += text
            
            # 使用pdfplumber提取表格
            tables = self._extract_tables(source)
            if tables:
                metadata["tables"] = tables
            
            # 构建文档结构
            sections = self._detect_api_sections(full_text)
            structure = DocumentStructure(sections=sections)
            
            return Document(
                content=full_text.strip(),
                metadata=metadata,
                structure=structure,
                file_type="pdf",
                file_path=str(file_path) if file_path else None
            )
        finally:
            doc.close()
    
    def _extract_tables(self, source: Union[Path, bytes, str]) -> List[Dict[str, Any]]:
        """提取PDF中的表格"""
        tables = []
        
        try:
            if isinstance(source, (str, Path)):
                pdf_file = pdfplumber.open(source)
            else:
                pdf_file = pdfplumber.open(io.BytesIO(source))
            
            for page_num, page in enumerate(pdf_file.pages):
                page_tables = page.extract_tables()
                for table in page_tables:
                    if table:
                        tables.append({
                            "page": page_num + 1,
                            "data": table
                        })
            
            pdf_file.close()
        except Exception as e:
            # 表格提取失败不影响整体加载
            pass
        
        return tables


class WordLoader(DocumentLoader):
    """Word文档加载器"""
    
    def load(self, source: Union[Path, bytes, str]) -> Document:
        """加载Word文档"""
        if isinstance(source, (str, Path)):
            file_path = Path(source)
            doc = DocxDocument(file_path)
        else:
            file_path = None
            doc = DocxDocument(io.BytesIO(source))
        
        full_text = ""
        sections = []
        tables = []
        
        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            full_text += text + "\n\n"
            
            # 检测标题样式
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                except ValueError:
                    level = 1
                sections.append(DocumentSection(
                    type=SectionType.HEADING,
                    content=text,
                    level=level
                ))
            else:
                sections.append(DocumentSection(
                    type=SectionType.PARAGRAPH,
                    content=text
                ))
        
        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                tables.append({
                    "index": table_idx,
                    "data": table_data
                })
                # 将表格内容也加入文本
                full_text += f"\n[Table {table_idx + 1}]\n"
                for row in table_data:
                    full_text += " | ".join(row) + "\n"
                full_text += "\n"
        
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "tables": tables
        }
        
        # 重新检测API相关章节
        sections = self._detect_api_sections(full_text)
        structure = DocumentStructure(sections=sections, tables=tables)
        
        return Document(
            content=full_text.strip(),
            metadata=metadata,
            structure=structure,
            file_type="docx",
            file_path=str(file_path) if file_path else None
        )


class ExcelLoader(DocumentLoader):
    """Excel文档加载器"""
    
    def load(self, source: Union[Path, bytes, str]) -> Document:
        """加载Excel文档"""
        if isinstance(source, (str, Path)):
            file_path = Path(source)
            xl_file = pd.ExcelFile(file_path)
        else:
            file_path = None
            xl_file = pd.ExcelFile(io.BytesIO(source))
        
        full_text = ""
        sheets_data = []
        
        for sheet_name in xl_file.sheet_names:
            df = xl_file.parse(sheet_name)
            
            # 将DataFrame转换为文本
            sheet_text = f"\n=== Sheet: {sheet_name} ===\n\n"
            sheet_text += df.to_string(index=False)
            sheet_text += "\n\n"
            
            full_text += sheet_text
            
            # 保存结构化数据
            sheets_data.append({
                "name": sheet_name,
                "data": df.to_dict(orient='records'),
                "columns": df.columns.tolist(),
                "row_count": len(df)
            })
        
        metadata = {
            "sheet_count": len(xl_file.sheet_names),
            "sheet_names": xl_file.sheet_names,
            "sheets": sheets_data
        }
        
        # Excel通常没有明显的API结构，但尝试检测
        sections = self._detect_api_sections(full_text)
        structure = DocumentStructure(sections=sections)
        
        return Document(
            content=full_text.strip(),
            metadata=metadata,
            structure=structure,
            file_type="xlsx",
            file_path=str(file_path) if file_path else None
        )


class TextLoader(DocumentLoader):
    """纯文本加载器"""
    
    def load(self, source: Union[Path, bytes, str]) -> Document:
        """加载纯文本"""
        if isinstance(source, (str, Path)):
            file_path = Path(source)
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        elif isinstance(source, bytes):
            file_path = None
            content = source.decode('utf-8')
        else:
            file_path = None
            content = source
        
        sections = self._detect_api_sections(content)
        structure = DocumentStructure(sections=sections)
        
        return Document(
            content=content.strip(),
            metadata={"char_count": len(content)},
            structure=structure,
            file_type="txt",
            file_path=str(file_path) if file_path else None
        )


def get_loader(file_type: str) -> DocumentLoader:
    """根据文件类型获取对应的加载器"""
    loaders = {
        "pdf": PDFLoader(),
        "docx": WordLoader(),
        "xlsx": ExcelLoader(),
        "txt": TextLoader(),
        "md": TextLoader(),
    }
    
    loader = loaders.get(file_type.lower())
    if not loader:
        raise ValueError(f"不支持的文件类型: {file_type}")
    
    return loader
